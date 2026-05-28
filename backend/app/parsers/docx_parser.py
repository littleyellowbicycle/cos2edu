import os
from app.parsers.base import BaseParser, ParseResult, ParseErrorCode


class DocxParser(BaseParser):
    supported_extensions = [".docx"]
    supported_mime_types = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ]

    async def parse(self, file_path: str) -> ParseResult:
        if not os.path.exists(file_path):
            return ParseResult(
                success=False,
                error_code=ParseErrorCode.UNKNOWN_ERROR,
                error_message=f"文件不存在: {file_path}",
            )

        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".doc":
            return ParseResult(
                success=False,
                error_code=ParseErrorCode.DOC_OLD_FORMAT,
                error_message="不支持旧版.doc格式，请转换为.docx格式",
                suggestion="可使用 Microsoft Word 或在线工具将.doc转换为.docx",
            )

        try:
            from docx import Document
        except ImportError:
            return ParseResult(
                success=False,
                error_code=ParseErrorCode.UNKNOWN_ERROR,
                error_message="python-docx 未安装，无法解析 DOCX 文件",
            )

        try:
            doc = Document(file_path)
            text_parts = []

            for para in doc.paragraphs:
                if para.text and para.text.strip():
                    text_parts.append(para.text.strip())

            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            content = "\n\n".join(text_parts)
            char_count = len(content)

            if char_count < 50:
                return ParseResult(
                    success=False,
                    error_code=ParseErrorCode.FILE_EMPTY,
                    error_message=f"DOCX文件内容为空或过少（仅{char_count}字）",
                    suggestion="请确认文件包含文字内容",
                )

            page_count = max(1, char_count // 2000)

            return ParseResult(
                success=True,
                text=content,
                page_count=page_count,
                char_count=char_count,
                format="docx",
            )

        except Exception as e:
            return ParseResult(
                success=False,
                error_code=ParseErrorCode.UNKNOWN_ERROR,
                error_message=f"DOCX解析失败: {str(e)}",
            )